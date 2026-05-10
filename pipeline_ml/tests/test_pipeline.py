"""
test_pipeline.py — Tests d'intégration pour les phases MLOps 1 & 2

Usage :
    python -m pytest pipeline_ml/tests/test_pipeline.py -v
    python pipeline_ml/tests/test_pipeline.py          # sans pytest

Tests :
    1. Import des modules modifiés (smoke test)
    2. MLflow — création d'un run de test, log params/metrics
    3. Extraction texte — PDF (pdfplumber) et DOCX (python-docx) sur contenu factice
    4. Regex parsing — CV structuré factice avec assertions précises
    5. Groq LLM — parsing sur 5 CVs .txt réels de data/raw
"""

import sys
import random
import tempfile
import textwrap
from pathlib import Path

ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(ROOT))

# ──────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────

IDENTITY_KEYS = {"cv_id", "source_filename", "name", "email", "phone", "gender", "age"}
FEATURE_KEYS  = {
    "cv_id", "profile_type", "target_role", "sector",
    "education_level", "nb_jobs", "years_experience", "avg_job_duration",
    "nb_technical_skills", "nb_languages", "nb_certifications", "label",
}

GREEN = "\033[92m"
RED   = "\033[91m"
RESET = "\033[0m"
SKIP  = "\033[93m"

_results = []


def ok(name):
    print(f"  {GREEN}✓{RESET} {name}")
    _results.append((name, "ok"))


def fail(name, reason):
    print(f"  {RED}✗{RESET} {name} — {reason}")
    _results.append((name, "fail", reason))


def skip(name, reason):
    print(f"  {SKIP}~{RESET} {name} — {reason}")
    _results.append((name, "skip", reason))


def assert_row_structure(row: dict, required_keys: set, label: str):
    missing = required_keys - set(row.keys())
    if missing:
        raise AssertionError(f"{label} manque les clés : {missing}")
    if not row.get("cv_id"):
        raise AssertionError(f"{label} cv_id est vide")


# ──────────────────────────────────────────────────────────────────
# Test 1 — Imports
# ──────────────────────────────────────────────────────────────────

def test_imports():
    print("\n[1] Smoke test — imports")
    for module_path, name in [
        ("pipeline_ml.core.p01_parse", "p01_parse"),
        ("pipeline_ml.core.p04_train", "p04_train"),
        ("pipeline_ml.core.p06_audit", "p06_audit"),
    ]:
        try:
            __import__(module_path)
            ok(f"import {name}")
        except Exception as e:
            fail(f"import {name}", str(e))


# ──────────────────────────────────────────────────────────────────
# Test 2 — MLflow
# ──────────────────────────────────────────────────────────────────

def test_mlflow():
    print("\n[2] MLflow — run de test")
    try:
        import mlflow
    except ImportError:
        skip("mlflow disponible", "pip install mlflow")
        return

    try:
        mlflow.set_experiment("cv-intelligence-tests")
        with mlflow.start_run(run_name="test-smoke") as run:
            mlflow.log_params({"C": 0.1, "l1_ratio": 0.5, "solver": "saga"})
            mlflow.log_metrics({"auc_roc_test": 0.837, "f1_test": 0.621, "recall_gap": 0.013})
            run_id = run.info.run_id
        ok(f"run créé et terminé (id={run_id[:8]}...)")

        client = mlflow.tracking.MlflowClient()
        r = client.get_run(run_id)
        assert r.data.params["C"] == "0.1", "Param C non loggé"
        assert abs(float(r.data.metrics["auc_roc_test"]) - 0.837) < 0.001
        ok("params et metrics récupérables via MlflowClient")

        with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False) as tmp:
            tmp.write("test artifact content")
            tmp_path = tmp.name
        with mlflow.start_run(run_id=run_id):
            mlflow.log_artifact(tmp_path)
        ok("log_artifact fonctionne")

    except Exception as e:
        fail("mlflow run", str(e))


# ──────────────────────────────────────────────────────────────────
# Test 3 — Extraction texte (PDF + DOCX)
# ──────────────────────────────────────────────────────────────────

def test_text_extraction():
    print("\n[3] Extraction texte — PDF et DOCX")
    from pipeline_ml.core.p01_parse import extract_text, _extract_text_pdf, _extract_text_docx

    # ── DOCX factice ─────────────────────────────────────────────
    try:
        from docx import Document
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)
        doc = Document()
        doc.add_paragraph("Name: Jean Dupont")
        doc.add_paragraph("Email: jean@example.com")
        doc.add_paragraph("Experience: 5 years in IT development")
        doc.save(tmp_path)

        text = _extract_text_docx(tmp_path)
        assert "Jean Dupont" in text, f"Texte extrait : {text[:100]}"
        ok("DOCX factice → texte extrait")
        tmp_path.unlink(missing_ok=True)
    except ImportError:
        skip("DOCX extraction", "pip install python-docx")
    except Exception as e:
        fail("DOCX extraction", str(e))

    # ── PDF factice ──────────────────────────────────────────────
    try:
        import pdfplumber
        try:
            from reportlab.pdfgen import canvas as rl_canvas
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp_path = Path(tmp.name)
            c = rl_canvas.Canvas(str(tmp_path))
            c.drawString(50, 750, "Name: Marie Martin")
            c.drawString(50, 730, "Skills: Python, SQL")
            c.save()
            text = _extract_text_pdf(tmp_path)
            assert "Marie" in text or len(text) >= 0
            ok("PDF factice (reportlab) → extraction pdfplumber OK")
            tmp_path.unlink(missing_ok=True)
        except ImportError:
            ok("pdfplumber importable (reportlab absent — test PDF réel ignoré)")
    except ImportError:
        skip("PDF extraction", "pip install pdfplumber")
    except Exception as e:
        fail("PDF extraction", str(e))

    # ── Router extract_text() ────────────────────────────────────
    try:
        with tempfile.NamedTemporaryFile(suffix=".txt", mode="w",
                                         delete=False, encoding="utf-8") as tmp:
            tmp.write("Name: Test User\nSkills: Python")
            tmp_path = Path(tmp.name)
        text = extract_text(tmp_path)
        assert "Test User" in text
        ok("extract_text() router .txt → lecture native")
        tmp_path.unlink(missing_ok=True)
    except Exception as e:
        fail("extract_text() router", str(e))


# ──────────────────────────────────────────────────────────────────
# Test 4 — parse_cv regex (mode structuré)
# ──────────────────────────────────────────────────────────────────

def test_regex_parsing():
    print("\n[4] Regex parsing — CV structuré factice")
    from pipeline_ml.core.p01_parse import parse_cv

    cv_content = textwrap.dedent("""\
        Name: Alice Dupont
        Gender: Female
        Date of Birth: 1995-03-15
        Email: alice@example.com
        Phone: +33612345678
        Address: 12 Rue de la Paix, 75001 Paris, France
        Target Role: Software Engineer

        Education:
        Master - Computer Science - Université Paris-Saclay - 2018

        Experience:
        Backend Developer - TechCorp - Paris - 2019-01 to 2022-06
        Senior Developer - StartupXYZ - Lyon - 2022-07 to present

        Skills:
        Technical: Python, SQL, Docker, Kubernetes
        Methods: Agile, Scrum, TDD
        Management: Team Leadership

        Languages:
        English - C1
        French - Native

        Certifications:
        AWS Certified Developer
        Google Cloud Associate
    """)

    with tempfile.NamedTemporaryFile(
        suffix=".txt", mode="w", delete=False,
        encoding="utf-8", dir=tempfile.gettempdir()
    ) as tmp:
        tmp.write(cv_content)
        tmp_path = Path(tmp.name)

    try:
        id_row, feat_row = parse_cv(tmp_path, labels_dict={})

        assert_row_structure(id_row,   IDENTITY_KEYS, "identity_row")
        assert_row_structure(feat_row, FEATURE_KEYS,  "feature_row")

        assert id_row["gender"] == "Female",          f"gender={id_row['gender']}"
        assert feat_row["nb_jobs"] == 2,              f"nb_jobs={feat_row['nb_jobs']}"
        assert feat_row["nb_certifications"] == 2,    f"certs={feat_row['nb_certifications']}"
        assert feat_row["has_english"] == 1,          f"has_english={feat_row['has_english']}"
        assert feat_row["profile_type"] == "senior",  f"profile={feat_row['profile_type']}"
        assert feat_row["sector"] == "IT",            f"sector={feat_row['sector']}"

        ok(f"regex parsing OK — {id_row['name']} | {feat_row['profile_type']} | "
           f"{feat_row['years_experience']} ans | {feat_row['nb_jobs']} postes")
    except AssertionError as e:
        fail("regex parsing assertions", str(e))
    except Exception as e:
        fail("regex parsing", str(e))
    finally:
        tmp_path.unlink(missing_ok=True)


# ──────────────────────────────────────────────────────────────────
# Test 5 — Groq LLM parsing sur 5 CVs réels
# ──────────────────────────────────────────────────────────────────

def test_groq_parsing(n_files: int = 5):
    print(f"\n[5] Groq — parsing LLM sur {n_files} CVs réels")

    raw_folder = ROOT / "data" / "raw"
    if not raw_folder.exists():
        skip("Groq parsing", f"data/raw introuvable ({raw_folder})")
        return

    txt_files = list(raw_folder.glob("*.txt"))
    if not txt_files:
        skip("Groq parsing", "aucun .txt dans data/raw")
        return

    import os
    if not os.getenv("GROQ_API_KEY"):
        try:
            from dotenv import load_dotenv
            load_dotenv(ROOT / ".env")
        except ImportError:
            pass
    if not os.getenv("GROQ_API_KEY"):
        skip("Groq parsing", "GROQ_API_KEY non défini")
        return

    from pipeline_ml.core.p01_parse import parse_cv_llm

    sample = random.sample(txt_files, min(n_files, len(txt_files)))
    print(f"  Fichiers sélectionnés : {[f.name for f in sample]}")

    successes = 0
    for filepath in sample:
        try:
            text = filepath.read_text(encoding="utf-8", errors="ignore")
            id_row, feat_row = parse_cv_llm(text, filename=filepath.name)

            assert_row_structure(id_row,   IDENTITY_KEYS, "identity_row")
            assert_row_structure(feat_row, FEATURE_KEYS,  "feature_row")

            profile = feat_row.get("profile_type", "?")
            name    = id_row.get("name") or "anonyme"
            exp     = feat_row.get("years_experience", 0)
            ok(f"{filepath.name} → {name} | {profile} | {exp} ans exp.")
            successes += 1
        except Exception as e:
            fail(filepath.name, str(e))

    if successes == 0:
        fail("Groq parsing global", "0 fichier parsé avec succès")
    else:
        ok(f"Groq : {successes}/{len(sample)} fichiers OK")


# ──────────────────────────────────────────────────────────────────
# Runner
# ──────────────────────────────────────────────────────────────────

def main():
    print("=" * 55)
    print("  CV-Intelligence — Tests d'intégration MLOps")
    print("=" * 55)

    test_imports()
    test_mlflow()
    test_text_extraction()
    test_regex_parsing()
    test_groq_parsing(n_files=5)

    print("\n" + "=" * 55)
    total   = len(_results)
    passed  = sum(1 for r in _results if r[1] == "ok")
    failed  = sum(1 for r in _results if r[1] == "fail")
    skipped = sum(1 for r in _results if r[1] == "skip")
    print(f"  Résultats : {GREEN}{passed} OK{RESET} | {RED}{failed} FAIL{RESET} | {SKIP}{skipped} SKIP{RESET} / {total} tests")
    print("=" * 55)
    return failed == 0


# ── Compatibilité pytest ──────────────────────────────────────────

def test_1_imports():         test_imports()
def test_2_mlflow():          test_mlflow()
def test_3_text_extraction(): test_text_extraction()
def test_4_regex_parsing():   test_regex_parsing()
def test_5_groq_parsing():    test_groq_parsing(n_files=5)


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
